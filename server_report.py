import docx
from docx.oxml.ns import qn # 한글 폰트
from docx.enum.text import WD_ALIGN_PARAGRAPH # 정렬
from docx.shared import Cm, Inches # 이미지 삽입시 길이 단위
import time # 파일생성 정렬방식
from docx2pdf import convert # docx를 pdf로 변환
## from fpdf import FPDF (keep)
## import pandas as pd (keep)

# 0. base
'''
* 사용시
from server_report import *

* 새로운 임포트 설치
pip3 install python-docx
pip3 install docx2pdf
pip3 install pypiwin32
pip3 install pywin32

* 윈도우에서 실행시 아나콘다에서 설치 (pywin32 - 윈도우)
conda install -c anaconda pywin32
conda update -n base conda
conda update --all
python -m pip install --upgrade pip
'''

# 1. word (미사용, 파이썬코드 만으로 docx생성 테스트)
'''
* 실행 예시 
server_report.word()
'''

# 2. word_form (기존 포맷에 수정하는 방식)
'''
* 전달 자료형
json형, ++ a

* 실행 예시
server_report.word_form(++a)

* 결과 예시
++ a
'''

# 3. convert_pdf
'''
* 실행 예시 
server_report.convert_pdf(1)

* 결과 예시
++ a

* 작동 원리
word_form에서 호출하는 식으로 작동 (word_form이 아닌 함수는 접근 X)
'''


class server_report:
    # 함수종류
    ### word : docx 생성 - 오직 코드로만 문서 생성 (미사용)
    ### word_form : 미리 정해진 양식(문서)을 사용하여 docx 생성
    ### convert_pdf : docx를 pdf로 변환
    ### ++ 타임랩스 만들기
    ### ++ 아이콘사용해서 가시화 하기

    # 전체적인 구조
    ### docx형식 파일을 불러와서 받은 데이터를 대입시킴
    ### docx를 pdf로 변환

    @staticmethod
    def word():
        # 공문서 : https://python-docx.readthedocs.io/en/latest/
        doc = docx.Document() # docx 생성

        para = doc.add_paragraph() 
        run = para.add_run("mini pot - title") # +본문
        run.font.size = docx.shared.Pt(30) # 폰트크기 
        run.bold = True # 볼트체 
        run.italic = True # 이텔릭체
        run.font.name = 'Cambria' # 폰트 설정
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조') # 한글 폰트 에러남
        last_paragraph = doc.paragraphs[-1]  # 이전 paragraph를 지정후 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 중앙정렬

        # 이미지 삽입
        doc.add_picture(".\\report_template\\test_image.png", width = Cm(13), height = Cm(8))

        # 단락 생성
        doc.add_paragraph('첫번째 단락', style='List Bullet')
        doc.add_paragraph('첫번째 순서 단락', style='List Number')

        # "test.docx" << 현제 위치에 생성
        doc.save(".\\report_result\\test.docx")

    @staticmethod
    def word_form(insert_data):
        # ++ 리눅스에서 폰트 같은 문제가 발생할 수 있기때문에 리눅스에서 테스트 해야함

        # report_form을 가져와서 새로운 docx를 만들고 저장할 꺼임!
        doc = docx.Document(".\\report_template\\report_form.docx")

        # 변수에 전달 받은 값을 넣어서 전달 - - - - - - - - - -
        # ++ insert_data의 값들을 매칭해 줘야함 지금은 임시로 그냥 준거임
        user = "ZIMyMeMine"
        plant = "T-hub"
        search_time = "nnnn/mm/dd" # ++ 하루, 일주일, 달 마다 다르게 처리해 줘야함...
        ## ++ NNNN년 MM월 DD일, NNNN년 MM월 4주, NNNN년 MM월
        # 식물사진2개, 온도-습도-영양액농도 그래프 사진 각각 1개
        plant_image_1 = ".\\report_template\\plant_image.jpg"
        plant_image_2 = ".\\report_template\\plant_image.jpg"
        temperature_image = ".\\report_template\\test_image.png"
        humidity_image = ".\\report_template\\test_image.png"
        nutrient_water_image = ".\\report_template\\test_image.png"
        ## ++ 이미지 경로에 대해서 전체적으로 설계가 필요함 
        lux = "all light : " + str(223)
        special = "comment"
        tip = "tip is tip"

        # 내용 기입 시작 - - - - - - - - - -
        para = doc.paragraphs

        # title
        para[0].text = user + "님의 " + plant + "성장보고서"
        ## doc.paragraphs[0].text = title // 한줄로 하면 이런코드 
        ## doc.add_heading("mini-pot", 0) // 단순 생성코드 
        # 조회구간
        para[1].text = "조회 구간 : " + search_time
        para[2].text = ""
        # 식물 이미지
        para[3].text = "# " + plant + " 상태"
        ## para[3].style = "contents_1" # !! 스타일 오류생기면 추가 아니면 패스
        para[4].text = ""
        # 빛총량
        para[12].text = lux
        # 특이사항
        para[15].text = special
        # 팁
        para[18].text = tip

        # 이미지 삽입 - - - - - - - - - -
        # 문서에서 테이블 정보 가져오기
        tables = doc.tables

        # 식물 이미지 1
        ### 0번째 테이블의 0행, 0열의 0칸에 기입, 중앙정렬
        para = tables[0].rows[0].cells[0].paragraphs[0]
        para.text = "< 시작 상태 >"
        ## para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ### 번째 테이블의 0행, 0열의 0칸에 이미지 추가 삽입 
        para = tables[0].rows[0].cells[0].paragraphs[0]
        run = para.add_run()
        run.add_picture(plant_image_1, width = Cm(7), height = Cm(5)) # !! 사진 비율 조정해줘야함
        ### ++ 사진의 위치 전달받기 or 디비에서 가져와서 전달받기
        ### ++ 하루일경우 이미지는 하나만 처리 할 수 있게끔  
        ## 3280 2464
        
        # 식물 이미지 2
        para = tables[0].rows[0].cells[1].paragraphs[0]
        para.text = "< 마지막 상태 >"
        para = tables[0].rows[0].cells[1].paragraphs[0]
        run = para.add_run()
        run.add_picture(plant_image_2, width = Cm(7), height = Cm(5)) # !! 사진 비율 조정해줘야함
        
        # 온도 그래프 이미지
        para = tables[1].rows[0].cells[0].paragraphs[0]
        para.text = ""
        ## para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run()
        run.add_picture(temperature_image, width = Cm(14), height = Cm(5)) # !! 사진 비율 조정해줘야함
        
        # 습도 그래프 이미지
        para = tables[2].rows[0].cells[0].paragraphs[0]
        para.text = ""
        run = para.add_run()
        run.add_picture(humidity_image, width = Cm(14), height = Cm(5)) # !! 사진 비율 조정해줘야함

        # 영양액농도 그래프 이미지
        para = tables[3].rows[0].cells[0].paragraphs[0]
        para.text = ""
        run = para.add_run()
        run.add_picture(nutrient_water_image, width = Cm(14), height = Cm(5)) # !! 사진 비율 조정해줘야함

        # 현제 봉인 << 어차피 고정일 태니까 인덱스 접근이 빠를꺼 같아서
        '''
        ## {}로 문자를 구별해서 접근 하는 방식(고정식으로 접근하면 오류가 생길까봐)
        for p in doc.paragraphs:

            if "{온도}" in p.text: 
                p.text = p.text.replace("{온도}","#test온")
                p.style = "contents_1"

            elif "{빛총량}" in p.text:
                p.add_run("light_all")
        '''

        # 마지막 문서쓰기 (문서 생성 시간) - - - - - - - - - -
        ## now_time = time.strftime('%Y%m%d-%H%M%S', time.localtime(time.time())) // 보기 안좋아서 바꿈
        time_local = time.localtime(time.time())
        time_Y = time.strftime('%Y', time_local)
        time_m = time.strftime('%m', time_local)
        time_d = time.strftime('%d', time_local)
        time_H = time.strftime('%H', time_local)
        time_M = time.strftime('%M', time_local)
        time_S = time.strftime('%S', time_local)
        now_time = time_Y + "/" + time_m + "/" + time_d + "-" + time_H + ":" + time_M + "." + time_S
        para = doc.add_paragraph("@ 발행 시간 : " + now_time, style = "last_text")

        # 디버깅 : word 문서를 번호 : 내용으로 확인 - - - - - - - - - -
        '''
        for x, paragraph in enumerate(doc.paragraphs):
            print(str(x) + " : " + paragraph.text)
        '''

        # 생성 날짜와 유저,식물 정보로 사진을 저장 - - - - - - - - - -
        ## ex) ZIMyMeMine_20210701-091731_T-hub
        ## ++ 지금은 하위 폴더(report_result)에 docx저장하고 pdf 변환 추후 경로 수정 필요
        now_time = time_Y + time_m + time_d + "-" + time_H + time_M + time_S
        file_path = ".\\report_result\\"
        file_name = user + "_" + now_time + "_" + plant
        doc.save(file_path + file_name + ".docx")

        # 생성된 파일이름을 pdf로 변환하기 위해 전달 - - - - - - - - - -
        server_report.convert_pdf(file_name)

    @staticmethod
    def convert_pdf(convert_file_name):
        
        # 받은 파일이름에 확장자랑 경로를 붙여줌
        docx = convert_file_name + ".docx"
        pdf = convert_file_name + ".pdf"
        docx_path = ".\\report_result\\" # !! ""안에 있으니까 \\ 이렇게 두번 해줘야함 
        pdf_path = ".\\report_result\\" # !! docx랑 pdf랑 동일 경로에 설정하면 안됨(윈도우) 리눅스는 아직 실험X but 하위 폴더일때는 가능

        # 변환
        convert(docx_path+docx, pdf_path+pdf)

        # 이부분에서 파일 경로를 전달할 send로 보내주면 됨
        ## ++ 상대경로를 줄지 아니면 파일 이름만 줘서 경로는 send에서 고정으로 할지는 send 생성후 다시
        print(pdf_path + pdf)

        
# test space ----------------------------------------
'''
server_report.word() # 단순 문서생성 테스트

<<<<<<< HEAD
#server_report.word_form(1) # 
=======
server_report.word_form(1)
>>>>>>> 8e7d4524058bf033716ea1c00715e6573f2fb84f

++이거 뭔 오류더라.... 있다가 todo에 등록하기...
'''

server_report.word_form(1)
