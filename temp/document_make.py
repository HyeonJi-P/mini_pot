import docx
from docx.oxml.ns import qn # 한글 폰트
from docx.enum.text import WD_ALIGN_PARAGRAPH # 중앙정렬
from docx.shared import Cm, Inches # 이미지 삽입시 길이 단위

import time # 파일생성시 이름에 날짜로 정렬

from docx2pdf import convert# docx를 pdf로 변환
import os
import sys
#import win32api

#from fpdf import FPDF
#import pandas as pd

# 0. base
'''
* 호출시 임포트
import docx

* 새로운 임포트 설치
pip3 install python-docx
pip3 install docx2pdf
pip3 install pypiwin32
pip3 install pywin32

conda install -c anaconda pywin32

'''

class document_make:
    # 함수종류
    ### word : docx 생성 (미사용)
    ### word_form : 미리 정해진 양식을 사용하여 docx 생성
    ### convert : docx를 pdf로 변환

    # 전체적인 구조
    ### docx형식 파일을 불러와서 받은 데이터를 대입시킴
    ### docx를 pdf로 변환

    @staticmethod
    def word(insert_data):
        doc = docx.Document()

        para = doc.add_paragraph()
        run = para.add_run('mini pot - title') # +본문
        run.font.size = docx.shared.Pt(30) # 폰트크기 
        run.bold = True # 볼트체 
        run.italic = True # 이텔릭체
        run.font.name = 'Cambria' # 폰트 설정
        # run._element.rPr.rFonts.set(qn('w:eastAsia'), '휴먼명조') # 한글 폰트 에러남
        last_paragraph = doc.paragraphs[-1]  # 이전 paragraph를 지정후 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 중앙정렬

        # 이미지 삽입
        ### ++ 이미지의 원래 비율에 맞개 줄이는 방법을 생각해야함
        ### ++ 아니면 규격을 정해놔서 거기에 맞춰야함
        doc.add_picture("test_image.png", width = Cm(13), height = Cm(8))

        # 단락 생성
        doc.add_paragraph('첫번째 단락', style='List Bullet')
        doc.add_paragraph('첫번째 순서 단락', style='List Number')

        # 공문서 : https://python-docx.readthedocs.io/en/latest/

        doc.save("time_plant.docx")

    @staticmethod
    def word_form(insert_data):
        print("START")
        doc = docx.Document("report_form.docx")

        title = "mini pot"
        user = "ZIMyMeMine"
        plant = "T-hub"

        now_time = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))

        # 제목 대입 
        ## doc.paragraphs[0].text = title // 한줄로
        ###doc.add_heading("mini-pot", 0) // 단순 생성
        para = doc.paragraphs
        para[0].text = title
        
        
        #doc.add_paragraph("ttap", style = 'a') // 스타일이 a인 내용 ttap를 생성

        #para = doc.add_paragraph() // 해당 라인의 뒤에 추가로 생성
        #run = para.add_run('mini pot - title')

        # ++ 보여주는 타입을 2개로 나누어서 일반:그래프, 고급:표



        # 디버깅 : word 문서를 번호 : 내용으로 확인
        for x, paragraph in enumerate(doc.paragraphs):
            print(str(x) + " : " + paragraph.text)

        # 생성 날짜와 유저,식물 정보로 사진을 저장 
        time_file = time.strftime('%Y%m%d-%H%M%S', time.localtime(time.time()))
        file_name = user + "_" + time_file + "_" + plant
        doc.save(file_name + ".docx")
        print("END")

        document_make.convert_pdf(file_name)

    @staticmethod
    def convert_pdf(convert_file_name):

        # !! 절대 경로시 \\ 이렇게 두번 해줘야함 
        
        docx = convert_file_name + ".docx"
        pdf = convert_file_name + ".pdf"
        docx_path = ".\\"
        pdf_path = ".\\"

        print(docx_path+docx)
        print(pdf_path+pdf)

        convert(docx_path+docx, pdf_path+pdf)

        

# test space ----------------------------------------
''''''
document_make.word_form(1)